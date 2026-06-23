"""Generate concise leadership progress speech as Word document on desktop."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

DESKTOP = Path.home() / "Desktop"
OUTPUT = DESKTOP / "一人团队工作台-领导汇报演讲稿.docx"


def set_run_font(run, size: int = 12, bold: bool = False, color: RGBColor | None = None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=20, bold=True, color=RGBColor(0x1A, 0x36, 0x5D))


def add_subtitle(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=10, color=RGBColor(0x88, 0x88, 0x88))


def add_section(doc: Document, title: str):
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run_font(run, size=13, bold=True, color=RGBColor(0x2E, 0x5C, 0x8A))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def add_line(doc: Document, text: str, bold_prefix: str = ""):
    p = doc.add_paragraph()
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=12, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=12)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(6)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)
    section.right_margin = Pt(90)

    add_title(doc, "一人团队工作台 · 阶段进展汇报")
    add_subtitle(doc, "建议时长 3 分钟")
    doc.add_paragraph()

    add_section(doc, "本阶段一句话")
    add_line(
        doc,
        "从「单端员工工具」升级为「一线展业 + 后台支持」双端协同平台，"
        "前后台支持链路在本版原型中首次跑通。",
    )

    add_section(doc, "三项核心进展")
    add_line(doc, "产品架构由单端扩展为双端——业务团队端服务一线展业，业务支持中心服务管理人员与中后台；登录按角色分流，两端并列、能力对齐。", "1. ")
    add_line(doc, "AI 能力从「个人助手」延伸为「条线助理体系」——业务团队端新增交叉验证助手，展业链路补全「分析—挖掘—方案—验证—服务」闭环；支持中心上线 7 个条线业务助理，任务角标、待办汇总、助理切换均已可用。", "2. ")
    add_line(doc, "协同机制从「各干各的」变为「一事一助理、全程可追溯」——一线发起需求，后台以对应条线助理身份接单回复；任务、对话、身份三端一致，支持响应可管可控。", "3. ")

    add_section(doc, "两端场景（各一句）")
    add_line(doc, "一线员工：打开即见今日任务与业绩，现场调 AI 团队完成客户分析与方案准备。", "业务团队 · ")
    add_line(doc, "管理人员/后台：按条线查看积压任务，切换助理身份处理一线请求，跨部门协同不再靠群聊接力。", "支持中心 · ")

    add_section(doc, "下步重点")
    add_line(doc, "选 1—2 条试点业务线，验证「一线发起—后台响应—结果回传」高频场景的实际效率。", "· ")
    add_line(doc, "对接真实任务与业绩数据源，从演示原型推进至可试用。", "· ")

    add_section(doc, "收尾")
    add_line(doc, "本版 Demo 已上线，可随时演示。以上汇报，请领导指示。")

    return doc


def main():
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
